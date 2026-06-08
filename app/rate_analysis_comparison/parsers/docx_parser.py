import os
import re
import logging
from pathlib import Path
from typing import Callable, Optional
import docx

from .base_parser import BaseParser, ParsedResult, ParseError, ParseWarning, RawRateRow, RawClauseChunk
from .pdf_parser import PdfParser, CONTEXT_PATTERN

logger = logging.getLogger(__name__)

class DocxParser(BaseParser):
    """
    Parser for Word (.docx) carrier rate agreements.
    Converts Word tables to grids and delegates parsing logic to PdfParser's engine.
    """

    def parse(self, file_path: str, progress_callback: Optional[Callable[[int, str], None]] = None) -> ParsedResult:
        if progress_callback:
            progress_callback(10, "Opening Word document...")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            doc = docx.Document(file_path)
        except Exception as e:
            raise ParseError(file_path, f"Failed to open DOCX file: {str(e)}")

        rate_rows: list[RawRateRow] = []
        clause_chunks: list[RawClauseChunk] = []
        warnings: list[ParseWarning] = []
        
        pdf_engine = PdfParser()
        page_context = {}

        # Scan the first 15 paragraphs for context/metadata (Carrier, Origin, Effective, etc.)
        for p in doc.paragraphs[:15]:
            text = p.text.strip()
            if text:
                m = CONTEXT_PATTERN.match(text)
                if m:
                    key = m.group(1).lower().strip()
                    val = m.group(2).strip()
                    page_context[key] = val

        num_tables = len(doc.tables)
        if num_tables == 0:
            # Check if there is any text to issue a Warning vs ParseError
            all_text = "\n".join(p.text for p in doc.paragraphs)
            form_signals = [
                "cover page", "submission", "carrier name", "signature",
                "please complete", "please confirm", "return submission",
                "primary contact", "billing assumption",
            ]
            form_hits = sum(1 for sig in form_signals if sig in all_text.lower())
            if form_hits >= 2:
                warnings.append(ParseWarning(
                    file_path, None,
                    "This Word document appears to be a cover page or form submission, not a rate "
                    "agreement. No rate data was extracted. Upload the correct file instead."
                ))
                return ParsedResult(
                    source_file=file_path,
                    rate_rows=[],
                    clause_chunks=[],
                    warnings=warnings,
                    sheet_notes=[all_text[:500]],
                    carrier_name=None,
                    effective_from=None,
                    effective_to=None,
                    service_level=None,
                    rates=[],
                    clauses=[]
                )
            
            raise ParseError(
                file_path,
                "Word document parsed successfully but contains no tables."
            )

        for i, table in enumerate(doc.tables):
            if progress_callback:
                pct = int(10 + (i / num_tables) * 80)
                progress_callback(pct, f"Parsing table {i+1} of {num_tables}...")

            # Convert table to 2D grid: list of list of str
            grid: list[list[str]] = []
            for row in table.rows:
                grid_row = []
                for cell in row.cells:
                    grid_row.append(cell.text.strip())
                grid.append(grid_row)

            if not grid or len(grid) < 2:
                continue

            # Process using PdfParser's Mode A table processor
            try:
                rows, chunks, warns = pdf_engine._process_raw_table(
                    grid, file_path, i + 1, page_context
                )
                rate_rows.extend(rows)
                clause_chunks.extend(chunks)
                warnings.extend(warns)
            except Exception as e:
                logger.warning("Failed to parse table %d: %s", i+1, e)

        if not rate_rows and not clause_chunks:
            # Check for cover page/form signal in paragraphs
            all_text = "\n".join(p.text for p in doc.paragraphs)
            form_signals = [
                "cover page", "submission", "carrier name", "signature",
                "please complete", "please confirm", "return submission",
                "primary contact", "billing assumption",
            ]
            form_hits = sum(1 for sig in form_signals if sig in all_text.lower())
            if form_hits >= 2:
                warnings.append(ParseWarning(
                    file_path, None,
                    "This Word document appears to be a cover page or form submission, not a rate "
                    "agreement. No rate data was extracted."
                ))
                return ParsedResult(
                    source_file=file_path,
                    rate_rows=[],
                    clause_chunks=[],
                    warnings=warnings,
                    sheet_notes=[all_text[:500]],
                    carrier_name=None,
                    effective_from=None,
                    effective_to=None,
                    service_level=None,
                    rates=[],
                    clauses=[]
                )
            raise ParseError(
                file_path,
                "Parser completed but found no rate rows and no clause data in the Word document."
            )

        if progress_callback:
            progress_callback(95, "Mapping schemas...")

        carrier_name, effective_from, effective_to, service_level, rates, clauses = pdf_engine._build_canonical_agreement(
            file_path, rate_rows, clause_chunks, page_context
        )

        return ParsedResult(
            source_file=file_path,
            rate_rows=rate_rows,
            clause_chunks=clause_chunks,
            warnings=warnings,
            sheet_notes=[p.text for p in doc.paragraphs[:10] if p.text],
            carrier_name=carrier_name,
            effective_from=effective_from,
            effective_to=effective_to,
            service_level=service_level,
            rates=rates,
            clauses=clauses,
        )
